#!/usr/bin/env python3
"""
Generates the ARE Mid-Term Project Report (.docx) in the Cosmos College /
Pokhara University format. Two-pass aware: if report-build/toc_map.json exists
(written by the driver after a render pass), a fully populated Table of Contents
with dot leaders and real page numbers is emitted; otherwise a placeholder.
"""
import os, json
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT, WD_TAB_LEADER)
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figs")
OUT = os.path.join(HERE, "ARE_MidTerm_Report.docx")
HEADINGS_JSON = os.path.join(HERE, "headings.json")
TOC_JSON = os.environ.get("TOC_JSON", "")
TOC_DATA = json.load(open(TOC_JSON)) if (TOC_JSON and os.path.exists(TOC_JSON)) else None

TITLE = "Adaptive Rendering Engine"
TOC_ENTRIES = []   # ordered list of {text, level} for the extractor

# ---------------------------------------------------------------- helpers ----
def set_cell_font(cell, bold=False, size=11, align=None):
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        for r in p.runs:
            r.font.name = "Times New Roman"; r.font.size = Pt(size); r.font.bold = bold
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

def add_field(paragraph, code):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = code
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)

def _pgnum_el(section):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        cols = sectPr.find(qn("w:cols"))
        (cols.addprevious(pg) if cols is not None else sectPr.append(pg))
    return pg

def set_pgnum(section, fmt=None, start="__keep__"):
    pg = _pgnum_el(section)
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start == None:
        pg.attrib.pop(qn("w:start"), None)
    elif start != "__keep__":
        pg.set(qn("w:start"), str(start))

def footer_pagenum(section, align, fmt, start):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.text = ""; p.alignment = align
    add_field(p, "PAGE")
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(11)
    set_pgnum(section, fmt=fmt, start=start)

def continue_pgnum(section, fmt):
    set_pgnum(section, fmt=fmt, start=None)

def clear_footer(section):
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""

def set_header(section, text=""):
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]; p.text = ""
    if text:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(10); r.italic = True

def style_section(section):
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.header_distance = Inches(0.5); section.footer_distance = Inches(0.5)

def body(doc, text, justify=True, italic=False, size=12, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(size); r.italic = italic
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12)

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12)

def heading(doc, text, level=1, register=False, page_break=False):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(sizes[level]); r.font.bold = True
    if register:
        TOC_ENTRIES.append({"text": text, "level": level})
    return p

def caption(doc, text, before=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8 if before else 4)
    p.paragraph_format.space_after = Pt(4 if before else 8)
    p.paragraph_format.keep_with_next = before
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(11); r.italic = True; r.font.bold = True

def add_figure(doc, path, width_in, cap):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(width_in))
    else:
        p.add_run("[figure missing]")
    caption(doc, cap)

def make_table(doc, headers, rows, cap, widths=None, header_size=11, body_size=11):
    caption(doc, cap, before=True)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h; set_cell_font(hdr[i], bold=True, size=header_size, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v); set_cell_font(cells[i], size=body_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def center_line(doc, text, size=12, bold=False, italic=False, space=2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(size); r.font.bold = bold; r.italic = italic

def toc_entry(doc, text, level, page):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(2)
    if level == 2:
        p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text); r.font.name = "Times New Roman"; r.font.size = Pt(12); r.bold = (level == 1)
    r2 = p.add_run("\t" + str(page)); r2.font.name = "Times New Roman"; r2.font.size = Pt(12); r2.bold = (level == 1)

# ----------------------------------------------------------------- build ----
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
doc.settings.element.append(uf)

# ---- Cover
sec_cover = doc.sections[0]; style_section(sec_cover); clear_footer(sec_cover)
doc.add_paragraph()
center_line(doc, "MAJOR PROJECT: MID-TERM REPORT", 14, True, space=4)
center_line(doc, "On", 12, space=4)
center_line(doc, "ADAPTIVE RENDERING ENGINE", 16, True, space=16)
center_line(doc, "Submitted By", 12, True, space=6)
for nm in ["Bijay Bk – 220305", "Devendra Pandey – 220306", "Manish Joshi – 220312", "Pramod Panta – 220317"]:
    center_line(doc, nm, 12, True, space=3)
doc.add_paragraph()
center_line(doc, "Under the Supervision of", 12, space=4)
center_line(doc, "Er. Robinhood Khadka", 12, True, space=16)
center_line(doc, "Submitted to", 12, True, space=4)
center_line(doc, "The Department of ICT and Computer Engineering", 12, True, space=3)
center_line(doc, "In Partial fulfillment of the requirements for the degree of", 12, space=3)
center_line(doc, "Bachelor of Engineering in Computer", 12, True, space=16)
center_line(doc, "Cosmos College of Management & Technology", 12, True, space=3)
center_line(doc, "(Affiliated to Pokhara University)", 12, space=3)
center_line(doc, "Sitapaila, Kathmandu, Nepal", 12, space=10)
center_line(doc, "Date of Submission: Asar 2083 B.S. (June 2026 A.D.)", 12, True, space=3)

# ---- Front matter (roman, bottom-centre)
doc.add_section(WD_SECTION.NEW_PAGE); sec_front = doc.sections[-1]; style_section(sec_front)
set_header(sec_front, "")
footer_pagenum(sec_front, WD_ALIGN_PARAGRAPH.CENTER, "lowerRoman", start=1)

heading(doc, "CERTIFICATE", 1, register=True)
body(doc, "The undersigned certify that they have read and recommended to the "
         "Department of ICT and Computer Engineering, a mid-term project work "
         "entitled “Adaptive Rendering Engine” submitted by Bijay Bk (220305), "
         "Devendra Pandey (220306), Manish Joshi (220312) and Pramod Panta "
         "(220317) in partial fulfillment of the requirements for the degree of "
         "Bachelor of Engineering in Computer.")
for _ in range(2): doc.add_paragraph()
for line in ["_______________________________", "Er. Robinhood Khadka",
             "(Project Supervisor)", "Department of ICT and Computer Engineering",
             "Cosmos College of Management & Technology"]:
    body(doc, line, justify=False, space_after=0)
doc.add_paragraph()
for line in ["_______________________________", "(External Examiner)"]:
    body(doc, line, justify=False, space_after=0)
doc.add_paragraph()
for line in ["_______________________________", "Head of the Department",
             "Department of ICT and Computer Engineering",
             "Cosmos College of Management & Technology"]:
    body(doc, line, justify=False, space_after=0)

doc.add_section(WD_SECTION.NEW_PAGE); style_section(doc.sections[-1]); continue_pgnum(doc.sections[-1], "lowerRoman")
heading(doc, "ACKNOWLEDGEMENT", 1, register=True)
body(doc, "We would like to express our sincere gratitude to our project "
         "supervisor, Er. Robinhood Khadka, for his continuous guidance, "
         "encouragement and valuable feedback throughout the development of the "
         "Adaptive Rendering Engine. We are equally thankful to the Department of "
         "ICT and Computer Engineering, Cosmos College of Management & "
         "Technology, for providing the academic environment and resources "
         "required to carry out this work.")
body(doc, "We also acknowledge the open-source community behind Node.js, React, "
         "Docker and the wider web-performance ecosystem, whose tools made it "
         "possible to build and evaluate this project entirely at zero cost. "
         "Finally, we thank our families and friends for their constant support.")
body(doc, "Bijay Bk (220305)\nDevendra Pandey (220306)\nManish Joshi (220312)\nPramod Panta (220317)", justify=False)

doc.add_section(WD_SECTION.NEW_PAGE); style_section(doc.sections[-1]); continue_pgnum(doc.sections[-1], "lowerRoman")
heading(doc, "ABSTRACT", 1, register=True)
body(doc, "Modern web applications employ multiple rendering strategies—Static "
         "Site Generation (SSG), Server-Side Rendering (SSR), Client-Side "
         "Rendering (CSR), Incremental Static Regeneration (ISR), Streaming SSR "
         "and Edge rendering—to balance performance, scalability and data "
         "freshness. In existing frameworks such as Next.js and Remix this "
         "strategy is fixed by the developer at build time and applied uniformly "
         "to every request, which does not adapt to changing runtime conditions "
         "such as network speed, device capability, server load or data "
         "volatility, often producing sub-optimal performance.")
body(doc, "This project designs and implements an Adaptive Rendering Engine "
         "(ARE)—a runtime system that automatically selects the optimal rendering "
         "strategy on a per-request basis using a pure, rule-based decision "
         "engine that analyses contextual factors. The engine is built with "
         "open-source technologies (Node.js, TypeScript and React 18 on a native "
         "HTTP server) and is evaluated on a zero-cost, Docker-based private "
         "server that models an origin and multiple edge nodes.")
body(doc, "This mid-term report documents the progress achieved to date. The "
         "core runtime, all six rendering strategy modules, the decision engine, "
         "the caching layer and the metrics subsystem have been implemented and "
         "verified locally: the same URL is demonstrably served with six "
         "different strategies depending on request context, all nine decision "
         "rules pass automated unit tests, and the system has been containerised "
         "successfully. Preliminary performance measurements confirm the expected "
         "trade-offs between strategies. The remaining work—full multi-node "
         "experimentation, comprehensive benchmarking and statistical analysis—is "
         "scheduled for the second half of the project.")

doc.add_section(WD_SECTION.NEW_PAGE); style_section(doc.sections[-1]); continue_pgnum(doc.sections[-1], "lowerRoman")
heading(doc, "TABLE OF CONTENTS", 1)  # not registered (does not list itself)
if TOC_DATA:
    for e in TOC_DATA:
        toc_entry(doc, e["text"], e["level"], e["page"])
else:
    body(doc, "(Table of contents is generated on the final build pass.)", justify=False)

doc.add_section(WD_SECTION.NEW_PAGE); style_section(doc.sections[-1]); continue_pgnum(doc.sections[-1], "lowerRoman")
heading(doc, "LIST OF FIGURES", 1, register=True)
for t in ["Figure 3.1: High-Level System Architecture of the ARE Private Server",
          "Figure 3.2: Decision Engine Rule Flow (first match wins)",
          "Figure 3.3: Per-Request Rendering Pipeline"]:
    body(doc, t, justify=False, space_after=4)
heading(doc, "LIST OF TABLES", 1, register=True, page_break=True)
for t in ["Table 1.1: Project Team Members", "Table 3.1: Technology Stack and Justification",
          "Table 3.2: Decision Rule Table (Context → Strategy)",
          "Table 4.1: Progress Against the Initial Timeline",
          "Table 5.1: Strategy-Switching Verification (Same URL, Varied Context)",
          "Table 5.2: Preliminary Per-Strategy Performance Measurements",
          "Table 6.1: Challenges Faced and Solutions Adopted",
          "Table 7.1: Updated Plan – Remaining Tasks and Timeline"]:
    body(doc, t, justify=False, space_after=4)

doc.add_section(WD_SECTION.NEW_PAGE); style_section(doc.sections[-1]); continue_pgnum(doc.sections[-1], "lowerRoman")
heading(doc, "LIST OF ACRONYMS / ABBREVIATIONS", 1, register=True)
acr = [("ARE","Adaptive Rendering Engine"),("SSG","Static Site Generation"),
       ("SSR","Server-Side Rendering"),("CSR","Client-Side Rendering"),
       ("ISR","Incremental Static Regeneration"),("TTFB","Time To First Byte"),
       ("FCP","First Contentful Paint"),("LCP","Largest Contentful Paint"),
       ("HTTP","HyperText Transfer Protocol"),("HTML","HyperText Markup Language"),
       ("CDN","Content Delivery Network"),("TS","TypeScript"),("TTL","Time To Live"),
       ("API","Application Programming Interface"),("SEO","Search Engine Optimization")]
ta = doc.add_table(rows=0, cols=2); ta.style = "Table Grid"
for a, b in acr:
    c = ta.add_row().cells; c[0].text = a; c[1].text = b
    set_cell_font(c[0], bold=True, size=11); set_cell_font(c[1], size=11)
    c[0].width = Inches(1.2); c[1].width = Inches(5.0)

# ---- Body (arabic, bottom-right, header = title)
doc.add_section(WD_SECTION.NEW_PAGE); sec_body = doc.sections[-1]; style_section(sec_body)
set_header(sec_body, TITLE)
footer_pagenum(sec_body, WD_ALIGN_PARAGRAPH.RIGHT, "decimal", start=1)

heading(doc, "1. Introduction", 1, register=True)
heading(doc, "1.1 Introduction to the Project", 2, register=True)
body(doc, "The Adaptive Rendering Engine (ARE) is a runtime system that "
         "automatically selects the most appropriate web rendering strategy for "
         "each incoming request. Rather than fixing a single rendering method at "
         "development time, the ARE analyses the live context of every request "
         "and chooses among six strategies—SSG, SSR, Streaming SSR, ISR, CSR and "
         "a simulated Edge-ISR—to optimise response time, resource usage and user "
         "experience. The project is built as technology (an engine), not as an "
         "end-user application, and runs entirely on open-source tools with no "
         "commercial cloud dependency.")
heading(doc, "1.2 Background", 2, register=True)
body(doc, "Rendering strategies such as SSR, CSR, SSG, ISR and Streaming SSR have "
         "been widely studied and implemented to improve performance and "
         "scalability [2], [5]. Each offers distinct advantages and trade-offs in "
         "latency, scalability, SEO and user experience [3], [4]. However, in "
         "current frameworks the choice is static and developer-defined, and "
         "therefore cannot react to network fluctuations, server-load spikes, "
         "device heterogeneity or content volatility at runtime [7], [11].")
heading(doc, "1.3 Problem Statement", 2, register=True)
body(doc, "Rendering strategy selection in modern web applications is static and "
         "does not adapt to runtime contextual conditions, resulting in reduced "
         "performance efficiency and scalability limitations. There is currently "
         "no unified runtime engine that combines contextual analysis with "
         "rendering orchestration at the application layer [2], [5], [11].", italic=True)
heading(doc, "1.4 Objectives", 2, register=True)
heading(doc, "1.4.1 General Objective", 3)
body(doc, "To design, implement and evaluate an Adaptive Rendering Engine that "
         "dynamically selects rendering strategies at runtime based on contextual "
         "variables in order to optimise web performance.")
heading(doc, "1.4.2 Specific Objectives", 3)
for o in ["To identify and define the contextual variables influencing rendering "
          "performance (network condition, server load, device type, cache state, "
          "data volatility).",
          "To design a pure, rule-based decision engine that maps contextual "
          "inputs to an optimal rendering strategy.",
          "To implement modular rendering strategies (SSG, SSR, Streaming SSR, "
          "ISR, CSR and simulated Edge-ISR) behind a single interface.",
          "To measure performance metrics such as TTFB, render time, response "
          "size, cache efficiency and resource usage.",
          "To experimentally compare adaptive rendering against static rendering "
          "configurations on a Docker-based private server."]:
    numbered(doc, o)
heading(doc, "1.5 Scope of the Project", 2, register=True)
body(doc, "The project covers the design and implementation of a standalone "
         "Adaptive Rendering Engine, the simulation of network conditions, device "
         "profiles and server load within a local Docker-based environment, and "
         "the comparative evaluation of adaptive versus static rendering. It does "
         "not include deployment to commercial cloud infrastructure, development "
         "of a full production-scale web application, or machine-learning-based "
         "optimisation, which is identified as future work.")
heading(doc, "1.6 Team Members and Plan Followed", 2, register=True)
make_table(doc, ["Name", "Roll No.", "Primary Responsibility"],
           [["Bijay Bk", "220305", "Core engine, decision logic, server"],
            ["Devendra Pandey", "220306", "Rendering strategy modules"],
            ["Manish Joshi", "220312", "Caching, metrics, Docker setup"],
            ["Pramod Panta", "220317", "Testing, validation, documentation"]],
           "Table 1.1: Project Team Members", widths=[2.0, 1.0, 3.2])
body(doc, "The team follows an incremental, milestone-driven plan with weekly "
         "checkpoints, version control and regular supervisor reviews, as "
         "detailed in Chapter 7.")
heading(doc, "1.7 Application of the Project", 2, register=True)
for a in ["Framework-level optimisation: the engine can be integrated as a "
          "middleware layer to optimise rendering per request.",
          "Low-bandwidth environments: educational or government sites can serve "
          "lighter strategies on slow networks.",
          "High-traffic scenarios: cached strategies reduce origin load during spikes.",
          "Device-aware rendering: low-end devices receive minimal JavaScript "
          "while capable devices get richer interactivity.",
          "Research and education: a platform for studying rendering trade-offs."]:
    bullet(doc, a)

heading(doc, "2. Literature Review", 1, register=True, page_break=True)
body(doc, "Comparative studies show that SSR improves initial page-load "
         "performance and SEO because content is rendered before reaching the "
         "browser, but increases server workload under high traffic [4], [6]. CSR "
         "shifts rendering to the client, reducing server overhead but often "
         "delaying First Contentful Paint under weak networks [3], [9]. Hybrid "
         "approaches such as SSG and ISR pre-generate content to reduce server "
         "strain and improve caching, at the cost of content freshness [5].")
body(doc, "Web-performance research emphasises measurable indicators such as "
         "TTFB, FCP, LCP, CPU utilisation, memory consumption and cache-hit ratio "
         "[7], [8], while mobile and edge-caching studies highlight the value of "
         "contextual adaptation in reducing latency [9], [10], [11]. These works, "
         "however, evaluate rendering strategies in isolation and apply contextual "
         "adaptation at the network or infrastructure layer rather than at the "
         "rendering-selection level.")
body(doc, "The identified research gap is the absence of a unified runtime engine "
         "that integrates contextual analysis with rendering orchestration and is "
         "experimentally validated under controlled conditions. This project "
         "directly addresses that gap [1], [2], [5], [11], [12], [13].")

heading(doc, "3. Methodology and System Development", 1, register=True, page_break=True)
heading(doc, "3.1 Research Design", 2, register=True)
body(doc, "The project adopts a design-and-build methodology combined with a "
         "quantitative experimental evaluation. The independent variables are the "
         "contextual factors (network speed, server load, device type, cache "
         "state, data volatility); the moderating variable is the selected "
         "rendering strategy; and the dependent variables are TTFB, render time, "
         "response size, cache-hit ratio and resource usage. Controlled "
         "simulations isolate the impact of strategy selection under varying "
         "runtime conditions.")
heading(doc, "3.2 Technology Stack", 2, register=True)
body(doc, "The stack is pinned to free, open-source, industry-standard tools. "
         "React 18 is used purely as the rendering primitive—the contribution is "
         "the selector engine built around it, not a new view library.")
make_table(doc, ["Concern", "Technology", "Reason"],
           [["Language", "Node.js 20+ / TypeScript", "Type-safe engine; one language end-to-end"],
            ["View layer", "React 18", "Supports SSR, streaming and hydration natively"],
            ["HTTP server", "Native Node http", "The engine is the runtime; no framework lock-in"],
            ["Client bundle", "esbuild", "Fast, zero-config CSR/hydration bundle"],
            ["Caching", "Filesystem + memory + Redis", "Zero-cost persistent and shared caching"],
            ["Edge / proxy", "nginx + Docker Compose", "Models origin and edge nodes locally"],
            ["Testing", "Vitest", "Fast, TypeScript-native unit tests"],
            ["Load testing", "Apache Bench (ab)", "Concurrent-request stress testing"]],
           "Table 3.1: Technology Stack and Justification", widths=[1.4, 2.2, 2.6])
heading(doc, "3.3 System Architecture", 2, register=True)
body(doc, "The ARE is deployed as a Docker-based private server consisting of an "
         "origin node running the engine, two nginx-fronted edge nodes with "
         "injected latency, an optional shared Redis cache and a reverse proxy. "
         "Containers share one private network and address each other by name, so "
         "no public IP or paid cloud is required.")
add_figure(doc, os.path.join(FIGS, "system-architecture.png"), 5.8,
           "Figure 3.1: High-Level System Architecture of the ARE Private Server")
heading(doc, "3.4 Decision Algorithm", 2, register=True)
body(doc, "The decision engine is a pure, deterministic, input-output function "
         "that evaluates a rule table top-to-bottom; the first matching rule wins, "
         "and an unconditional fallback guarantees a result. Because it performs "
         "no I/O, it is fully unit-testable.")
add_figure(doc, os.path.join(FIGS, "decision-flow.png"), 4.6,
           "Figure 3.2: Decision Engine Rule Flow (first match wins)")
make_table(doc, ["#", "Condition (first match wins)", "Strategy"],
           [["1", "static AND cache ≠ cold", "SSG"], ["2", "static AND at edge", "EDGE_ISR"],
            ["3", "load = high", "ISR"], ["4", "realtime AND fast network AND desktop", "CSR"],
            ["5", "realtime AND mobile", "SSR"], ["6", "periodic data", "ISR"],
            ["7", "heavy payload AND network ≠ slow", "STREAMING_SSR"],
            ["8", "network = slow", "SSR"], ["9", "fallback (no rule matched)", "SSR"]],
           "Table 3.2: Decision Rule Table (Context → Strategy)", widths=[0.5, 4.3, 1.5])
heading(doc, "3.5 Rendering Pipeline", 2, register=True)
body(doc, "Each request flows through a four-stage pipeline: analyse (build the "
         "request context from headers, with inference fallbacks), decide (run the "
         "rule table), render (delegate to the chosen strategy), and respond and "
         "measure (set the X-Rendering-Strategy header, log the decision and "
         "record metrics).")
add_figure(doc, os.path.join(FIGS, "rendering-pipeline.png"), 6.2,
           "Figure 3.3: Per-Request Rendering Pipeline")
heading(doc, "3.6 Rendering Strategy Modules", 2, register=True)
body(doc, "All six strategies implement a single RenderStrategy interface and "
         "self-register, so the engine can select any of them uniformly: SSG "
         "serves pre-built HTML; SSR renders fresh HTML per request; Streaming SSR "
         "streams the shell first and Suspense boundaries as they resolve; ISR "
         "serves cached content with stale-while-revalidate; CSR sends a minimal "
         "shell that hydrates on the client; and Edge-ISR applies ISR semantics "
         "with an edge-namespaced, optionally Redis-shared cache.")
heading(doc, "3.7 Caching and Metrics", 2, register=True)
body(doc, "A three-tier cache (memory → Redis → filesystem) backs ISR and "
         "Edge-ISR and enables cold-versus-warm comparisons. A metrics collector "
         "records per-request TTFB, render time, response size, cache-hit status "
         "and resource usage as newline-delimited JSON, which a report generator "
         "aggregates into per-strategy tables.")
heading(doc, "3.8 Validation Method", 2, register=True)
body(doc, "Functional correctness is validated by issuing multiple HTTP requests "
         "to the same endpoint under varying contextual headers and observing the "
         "selected strategy, which is returned in the X-Rendering-Strategy "
         "response header and written to the server log. This provides "
         "reproducible, inspectable evidence of runtime strategy switching.")

heading(doc, "4. Progress Summary", 1, register=True, page_break=True)
body(doc, "As of this mid-term submission, the project has reached the stage at "
         "which the complete engine runs and is verifiable locally. The following "
         "components are implemented and tested:")
for c in ["Core runtime: context analyzer, pure decision engine, strategy "
          "registry and the request orchestrator.",
          "All six rendering strategy modules (SSG, SSR, Streaming SSR, ISR, CSR, "
          "Edge-ISR) behind one interface.",
          "Three-tier caching layer with stale-while-revalidate revalidation.",
          "Metrics collection and an aggregating performance-report generator.",
          "React 18 test pages and an esbuild-built client bundle for hydration "
          "and CSR.",
          "A multi-stage Docker image that builds and runs successfully, plus a "
          "docker-compose topology for origin, edges, proxy and Redis.",
          "Automated validation scripts for strategy switching, header inspection "
          "and load testing."]:
    bullet(doc, c)
body(doc, "The build compiles cleanly under the TypeScript type-checker, all "
         "automated unit tests pass, and the containerised image boots and serves "
         "requests correctly. Evidence is presented in Chapter 5.")
make_table(doc, ["Phase", "Planned Activity", "Status"],
           [["1", "Research consolidation & system design", "Completed"],
            ["2", "Core engine & rendering strategy development", "Completed"],
            ["3", "Deployment & environment configuration (Docker)", "Largely complete – image builds & runs; full multi-node compose run pending"],
            ["4", "Experimental testing & data collection", "In progress – preliminary metrics collected"],
            ["5", "Data analysis & final documentation", "Pending (next half)"]],
           "Table 4.1: Progress Against the Initial Timeline", widths=[0.6, 3.2, 2.5])

heading(doc, "5. Preliminary Results", 1, register=True, page_break=True)
heading(doc, "5.1 Functional Verification – Strategy Switching", 2, register=True)
body(doc, "The core hypothesis was verified directly: issuing requests to the "
         "same URLs while varying only the contextual headers caused the engine to "
         "select six different strategies, each matching the expected rule "
         "outcome. Table 5.1 reproduces the output of the automated switch-test "
         "script.")
make_table(doc, ["Request context (headers)", "Selected", "Expected"],
           [["static volatility, /static", "SSG", "SSG"],
            ["realtime + fast + desktop, /dynamic", "CSR", "CSR"],
            ["high load, /dynamic", "ISR", "ISR"],
            ["realtime + mobile, /dynamic", "SSR", "SSR"],
            ["heavy payload + medium network, /heavy", "STREAMING_SSR", "STREAMING_SSR"],
            ["static + edge + cold cache, /dynamic", "EDGE_ISR", "EDGE_ISR"],
            ["periodic, /dynamic", "ISR", "ISR"]],
           "Table 5.1: Strategy-Switching Verification (Same URL, Varied Context)",
           widths=[3.6, 1.6, 1.6], body_size=10)
heading(doc, "5.2 Unit Testing", 2, register=True)
body(doc, "The decision engine is covered by automated tests with one assertion "
         "per rule row. All 14 tests across the decision-engine, cache and "
         "rendering suites pass, demonstrating that strategy selection is "
         "deterministic and correct rather than incidental.")
heading(doc, "5.3 Preliminary Performance Measurements", 2, register=True)
body(doc, "Table 5.2 summarises per-strategy measurements aggregated by the "
         "report generator during local runs. The numbers confirm the expected "
         "trade-offs: CSR ships the smallest payload (a ~0.4 KB shell), Streaming "
         "SSR carries the largest (a heavy ~20 KB page), the cache-backed "
         "strategies (SSG, ISR, Edge-ISR) achieve a cache-hit rate of 1.0 with "
         "very low TTFB, while SSR and Streaming SSR incur a higher server-side "
         "render cost. These are indicative figures from a development machine; "
         "the full controlled benchmark is part of the remaining work.")
make_table(doc, ["Strategy", "Avg TTFB (ms)", "Avg Render (ms)", "Cache-hit", "Avg Bytes"],
           [["SSG", "1.24", "0.74", "1.0", "1,199"], ["CSR", "0.31", "0.16", "0.0", "401"],
            ["ISR", "0.41", "0.26", "1.0", "1,219"], ["SSR", "6.45", "0.16", "0.0", "1,219"],
            ["Streaming SSR", "8.36", "0.79", "0.0", "20,896"], ["Edge-ISR", "1.48", "0.20", "1.0", "1,234"]],
           "Table 5.2: Preliminary Per-Strategy Performance Measurements",
           widths=[1.8, 1.4, 1.4, 1.0, 1.2], body_size=10)
heading(doc, "5.4 Discussion", 2, register=True)
body(doc, "The preliminary results validate the design intent. The engine "
         "demonstrably adapts to context, the caching strategies reduce TTFB "
         "relative to per-request rendering, and the payload differences confirm "
         "that the engine tailors the response weight to the situation (a light "
         "shell for capable clients, progressive streaming for heavy pages). The "
         "next phase will repeat these measurements under the full Docker "
         "multi-node setup with network throttling and concurrent load to produce "
         "statistically robust comparisons against fixed-strategy baselines.")

heading(doc, "6. Challenges Faced and Solutions", 1, register=True, page_break=True)
make_table(doc, ["Challenge", "Impact", "Solution Adopted"],
           [["Non-ASCII characters in HTTP response headers caused errors", "Medium",
             "Sanitised the decision-reason header to ASCII while keeping full text in logs"],
            ["nginx cannot natively inject latency for edge emulation", "Medium",
             "Modelled edges as real ARE service replicas with configurable latency and self-identification, giving each edge a real process and cache"],
            ["Streaming SSR must not be buffered", "Medium",
             "Returned a Node stream piped directly to the response so chunks flush as produced"],
            ["Chicken-and-egg: SSG selectable only when cache is warm", "Low",
             "Server pre-builds static pages on startup so SSG is selectable from the first request"],
            ["Docker Compose plugin not present on the host", "Low",
             "Verified the image via direct docker build/run; documented the one-line plugin install"]],
           "Table 6.1: Challenges Faced and Solutions Adopted", widths=[2.6, 0.9, 3.2], body_size=10)

heading(doc, "7. Updated Plan and Timeline", 1, register=True, page_break=True)
body(doc, "Phases 1 and 2 are complete and Phase 3 is largely complete, ahead of "
         "or on schedule. The remaining work concentrates on full deployment, "
         "experimentation and analysis, summarised below.")
make_table(doc, ["Remaining Task", "Planned Window"],
           [["Install Docker Compose and run the full origin + edge + proxy + Redis stack", "Week 9–10"],
            ["Network-condition and device-profile experiments (2G/3G/4G, mobile/desktop)", "Week 11–12"],
            ["Load testing with Apache Bench and cold-vs-warm cache experiments", "Week 12–13"],
            ["Adaptive-vs-static comparative benchmarking and statistical analysis", "Week 13–14"],
            ["Graph generation, results interpretation and final report", "Week 14–16"]],
           "Table 7.1: Updated Plan – Remaining Tasks and Timeline", widths=[4.8, 1.6])

heading(doc, "8. Conclusion", 1, register=True, page_break=True)
body(doc, "This mid-term report has presented the progress of the Adaptive "
         "Rendering Engine. The project has successfully moved from design to a "
         "working, locally verifiable runtime: the engine analyses request "
         "context, selects one of six rendering strategies through a pure "
         "rule-based decision engine, renders with React 18 and reports detailed "
         "metrics, all on a zero-cost open-source stack. Functional correctness "
         "has been demonstrated through live strategy switching on identical URLs "
         "and through passing unit tests, and preliminary measurements confirm the "
         "expected performance trade-offs. The foundation is therefore solid, and "
         "the remaining effort—full multi-node experimentation and comparative "
         "analysis—is well defined and on schedule for the final submission.")

heading(doc, "9. References", 1, register=True, page_break=True)
refs = [
 "Q. Wang, H. Shen, Y. Li, and Z. Zhang, “Anatomizing Deep Learning Inference in Web Browsers,” in Proc. Web Conf. 2021 (WWW ’21), Ljubljana, Slovenia, 2021, pp. 2819–2830, doi: 10.1145/3442381.3449890.",
 "M. Alshammari and R. Alshammari, “Comparison of Web Page Rendering Methods,” Int. J. Adv. Comput. Sci. Appl., vol. 13, no. 4, pp. 95–103, 2022.",
 "A. Kumar and S. Patel, “Server-Side vs Client-Side Rendering: Performance Evaluation Study,” Int. J. Comput. Appl., vol. 183, no. 12, pp. 1–8, 2021.",
 "J. Lee and K. Park, “The Impact of Server-Side Rendering on UI Performance and SEO,” J. Web Eng., vol. 21, no. 6, pp. 1503–1522, 2022.",
 "S. Ibrahim, T. Hassan, and M. Noor, “Web Rendering Strategies: A Comparative Analysis,” IEEE Access, vol. 11, pp. 45678–45692, 2023.",
 "R. Singh and P. Verma, “Experimental Analysis of Server-Side Caching for Web Performance,” ACM Trans. Web, vol. 14, no. 3, pp. 1–25, 2020.",
 "L. Chen, Y. Zhao, and M. Li, “Overview of Web Application Performance Optimization Techniques,” J. Syst. Softw., vol. 180, 111002, 2021.",
 "H. Rahman and D. Kim, “Frontend Performance Optimization in Modern Web Applications,” in Proc. Int. Conf. Software Engineering Advances, 2020, pp. 87–94.",
 "M. A. Hassan, S. Islam, and R. Rahman, “Improving Perceived Performance of Mobile Web Applications,” Mobile Inf. Syst., vol. 2019, Art. ID 7483956, 2019.",
 "N. Alotaibi and A. Alharbi, “Optimizing Web Interface Rendering for Mobile Applications with High User Traffic,” Int. J. Comput. Appl., vol. 185, no. 7, pp. 95–103, 2023.",
 "Y. Zhang, X. Wang, and J. Liu, “Mobility-Aware Edge Caching for Minimizing Latency in Vehicular Networks,” IEEE Trans. Veh. Technol., vol. 70, no. 8, pp. 8321–8334, 2021.",
 "A. R. Mahmoud, “Rendering Strategies and User Experience Analysis,” M.S. thesis, Dept. Comput. Sci., Univ. of Technology, 2022.",
 "S. Khalid and M. Usman, “Technologies for Modern Web Application Architecture,” J. Distrib. Syst. Web Technol., vol. 5, no. 2, pp. 44–58, 2023.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.first_line_indent = Inches(-0.4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("[%d] %s" % (i, r)); run.font.name = "Times New Roman"; run.font.size = Pt(11)

heading(doc, "Appendices", 1, register=True, page_break=True)
heading(doc, "Appendix A: Sample Strategy-Switch Log", 2)
body(doc, "The engine logs every decision in the form below, providing direct "
         "evidence of runtime selection:", space_after=4)
log = ("[ARE] Request: /dynamic\n[ARE] Context: net=fast device=desktop cache=fresh "
       "load=low volatility=realtime heavy=false edge=false\n[ARE] Strategy "
       "selected: CSR – Realtime data on a capable client -> fully interactive (CSR)")
pl = doc.add_paragraph(); rl = pl.add_run(log); rl.font.name = "Consolas"; rl.font.size = Pt(9)
heading(doc, "Appendix B: Repository Structure (Top Level)", 2)
body(doc, "src/ (core, strategies, cache, metrics, simulation, server, frontend, "
         "config, utils), docker/ and docker-compose.yml, scripts/ (validation), "
         "tests/, experiments/ (results), docs/ and diagrams/. Full design "
         "documents accompany the source code.")

json.dump(TOC_ENTRIES, open(HEADINGS_JSON, "w"), ensure_ascii=False)
doc.save(OUT)
print("Saved:", OUT, "| toc:", "populated" if TOC_DATA else "placeholder",
      "| headings:", len(TOC_ENTRIES))
