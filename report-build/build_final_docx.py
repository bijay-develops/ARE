#!/usr/bin/env python3
"""
Renders the ARE Final Project Report to .docx from the shared content model,
in the Cosmos College / Pokhara University BE project-report format:

  A4 portrait, Times New Roman, 12 pt body, 1.5 line spacing, justified,
  1 inch margins, 0.5 inch header/footer, lowercase roman page numbers
  (bottom centre) for front matter and arabic (bottom right) for the body,
  with the project title as a running header on body pages.

The table of contents is emitted as a real Word TOC field whose cached result
is pre-populated from the page map produced by the PDF build, so the entries
and page numbers are visible in any viewer and refresh exactly in Word.
"""
import os, sys, json

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import final_content as C

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,
                            WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK)
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

OUT = os.path.join(HERE, "ARE_Final_Report.docx")
TOC_PAGES = os.path.join(HERE, "data", "toc_pages.json")
TOCMAP = json.load(open(TOC_PAGES)) if os.path.exists(TOC_PAGES) else []

FONT = "Times New Roman"
MONO = "Courier New"
AVAIL_IN = 6.27          # A4 width minus two 1-inch margins

# ── low-level helpers ────────────────────────────────────────────────────────
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e

def run_font(r, size=12, bold=False, italic=False, name=FONT):
    r.font.name = name; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    r._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    return r

def add_field(paragraph, code):
    """Insert a simple field (e.g. PAGE) with no cached result."""
    run = paragraph.add_run()
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    it = _el("w:instrText", **{"xml:space": "preserve"}); it.text = code
    run._r.append(it)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))
    return run

def _pgnum(section):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = _el("w:pgNumType")
        cols = sectPr.find(qn("w:cols"))
        (cols.addprevious(pg) if cols is not None else sectPr.append(pg))
    return pg

def set_pgnum(section, fmt=None, start="keep"):
    pg = _pgnum(section)
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is None:
        pg.attrib.pop(qn("w:start"), None)
    elif start != "keep":
        pg.set(qn("w:start"), str(start))

def style_section(section):
    section.page_width = Mm(210); section.page_height = Mm(297)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Inches(1))
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

def footer_page(section, align, fmt, start):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.text = ""; p.alignment = align
    add_field(p, "PAGE")
    for r in p.runs:
        run_font(r, 11)
    set_pgnum(section, fmt=fmt, start=start)

def clear_footer(section):
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""

def set_header(section, text=""):
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]; p.text = ""
    if text:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_font(p.add_run(text), 10, italic=True)

def new_section(doc, fmt, start="keep", header="", align=WD_ALIGN_PARAGRAPH.CENTER):
    doc.add_section(WD_SECTION.NEW_PAGE)
    s = doc.sections[-1]
    style_section(s); set_header(s, header)
    footer_page(s, align, fmt, start)
    return s

# ── block writers ────────────────────────────────────────────────────────────
def para(doc, text, *, justify=True, italic=False, size=12, after=8, before=0,
         indent=0, spacing=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(after); pf.space_before = Pt(before)
    if indent:
        pf.left_indent = Inches(indent); pf.right_indent = Inches(indent)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    run_font(p.add_run(text), size, italic=italic)
    return p

def centre(doc, text, size=12, bold=False, italic=False, after=3):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    run_font(p.add_run(text), size, bold=bold, italic=italic)
    return p

HEAD_SIZE = {1: 16, 2: 14, 3: 12}

def heading(doc, text, level=1, page_break=False):
    """Real Word heading styles so the TOC field and navigation pane work."""
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    pf.space_after = Pt(6); pf.keep_with_next = True
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run_font(p.add_run(text), HEAD_SIZE[level], bold=True)
    return p

def caption(doc, text, before=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(9 if before else 5); pf.space_after = Pt(5 if before else 10)
    pf.keep_with_next = before
    run_font(p.add_run(text), 11, bold=True, italic=True)

def set_cell(cell, text, *, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(1.5); pf.space_after = Pt(1.5)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run_font(p.add_run(str(text)), size, bold=bold)

def shade(cell, hexcolor="E8E8E8"):
    cell._tc.get_or_add_tcPr().append(_el("w:shd", **{"w:val": "clear", "w:fill": hexcolor}))

def add_table(doc, blk):
    caption(doc, blk["cap"], before=True)
    headers, rows = blk["headers"], blk["rows"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    widths = blk["widths"]
    if widths:
        total = sum(widths)
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w / total * AVAIL_IN)
    # repeat the header row when the table spans pages
    trPr = t.rows[0]._tr.get_or_add_trPr()
    trPr.append(_el("w:tblHeader", **{"w:val": "true"}))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code(doc, text):
    lines = text.split("\n")
    longest = max((len(l) for l in lines), default=1)
    size = min(8.4, (AVAIL_IN * 72 - 16) / (0.6 * max(longest, 1)))
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]; cell.width = Inches(AVAIL_IN)
    shade(cell, "F4F4F4")
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run_font(p.add_run(line if line.strip() else "\u00a0"), size, name=MONO)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_figure(doc, blk):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    if os.path.exists(blk["path"]):
        p.add_run().add_picture(blk["path"], width=Inches(blk["width"]))
    else:
        run_font(p.add_run("[figure missing]"), 11, italic=True)
    caption(doc, blk["cap"])

def add_list(doc, items, numbered):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_font(p.add_run(it), 12)

# ── table of contents field with a pre-populated cached result ───────────────
def toc_entry_paragraph(doc, text, level, page):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(2)
    if level >= 1:
        pf.left_indent = Inches(0.32)
    pf.tab_stops.add_tab_stop(Inches(AVAIL_IN), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run_font(p.add_run(text), 12, bold=(level == 0))
    run_font(p.add_run("\t" + str(page)), 12, bold=(level == 0))
    return p

def add_toc(doc):
    """begin -> instrText -> separate -> [cached entries] -> end"""
    start = doc.add_paragraph()
    r = start.add_run()
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    it = _el("w:instrText", **{"xml:space": "preserve"})
    it.text = r'TOC \o "1-2" \h \z \u'
    r._r.append(it)
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    start.paragraph_format.space_after = Pt(0)

    for e in TOCMAP:
        toc_entry_paragraph(doc, e["text"], e["level"], e["page"])

    end = doc.add_paragraph()
    er = end.add_run()
    er._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))
    end.paragraph_format.space_after = Pt(0)

# ════════════════════════════════════════════════════════════════════ build ══
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
for lvl in (1, 2, 3):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = FONT; st.font.size = Pt(HEAD_SIZE[lvl])
    st.font.bold = True; st.font.color.rgb = RGBColor(0, 0, 0)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
doc.settings.element.append(_el("w:updateFields", **{"w:val": "true"}))

def title_page(with_supervisor):
    doc.add_paragraph()
    centre(doc, "AN ENGINEERING PROJECT REPORT", 14, True, after=6)
    centre(doc, "On", 12, after=6)
    centre(doc, C.TITLE_UPPER, 18, True, after=26)
    centre(doc, "Submitted By", 12, True, after=8)
    for nm, rl in C.AUTHORS:
        centre(doc, f"{nm} – {rl}", 12, True, after=4)
    doc.add_paragraph()
    if with_supervisor:
        centre(doc, "Under the Supervision of", 12, after=5)
        centre(doc, C.SUPERVISOR, 12, True, after=24)
    else:
        doc.add_paragraph()
    centre(doc, "Submitted to", 12, True, after=5)
    centre(doc, C.DEPARTMENT, 12, True, after=4)
    centre(doc, "In Partial fulfillment of the requirements for the degree of", 12, after=4)
    centre(doc, C.DEGREE, 12, True, after=26)
    centre(doc, C.COLLEGE, 12, True, after=4)
    centre(doc, C.AFFIL, 12, after=4)
    centre(doc, C.PLACE, 12, after=14)
    centre(doc, C.SUBMISSION, 12, True, after=3)

# ---- cover (unnumbered)
sec = doc.sections[0]; style_section(sec); clear_footer(sec); set_header(sec, "")
title_page(False)

# ---- front matter (roman, bottom centre)
new_section(doc, "lowerRoman", start=1)
title_page(True)

new_section(doc, "lowerRoman")
heading(doc, "COPYRIGHT", 1)
para(doc, "The author has agreed that the Library, Pokhara University, Cosmos College of "
          "Management & Technology, may make this engineering project report freely available "
          "for inspection. Moreover, the author has agreed that permission for extensive "
          "copying of this report for scholarly purpose may be granted by the supervisor who "
          "supervised the work recorded herein or, in their absence, by the college authority "
          "in which the project work was done. Copying or any other use of this report for "
          "financial gain without approval of the college and the author's permission is "
          "strictly prohibited.")
para(doc, "Request for the permission to copy or to make any other use of the materials in "
          "this report in whole or in part should be addressed to:")
doc.add_paragraph()
para(doc, C.COLLEGE, justify=False, after=0)
para(doc, C.PLACE, justify=False, after=0)

new_section(doc, "lowerRoman")
heading(doc, "CERTIFICATE", 1)
names = ", ".join(f"{n} ({r})" for n, r in C.AUTHORS[:-1])
para(doc, f"The undersigned certify that they have read and recommended to the Department of "
          f"ICT and Computer Engineering, a final year project work entitled “{C.TITLE}” "
          f"submitted by {names} and {C.AUTHORS[-1][0]} ({C.AUTHORS[-1][1]}) in partial "
          f"fulfillment of the requirements for the degree of {C.DEGREE}.")
for _ in range(2):
    doc.add_paragraph()
for grp in [["_______________________________", C.SUPERVISOR, "(Project Supervisor)",
             "Department of ICT and Computer Engineering", C.COLLEGE],
            ["_______________________________", "(External Examiner)"],
            ["_______________________________", "Head of the Department",
             "Department of ICT and Computer Engineering", C.COLLEGE]]:
    for line in grp:
        para(doc, line, justify=False, after=0)
    doc.add_paragraph()

new_section(doc, "lowerRoman")
heading(doc, "ACKNOWLEDGEMENT", 1)
for t in [
    "We would like to express our sincere gratitude to our project supervisor, "
    f"{C.SUPERVISOR}, for his continuous guidance, encouragement and valuable feedback "
    "throughout the design, implementation and evaluation of the Adaptive Rendering Engine. "
    "His insistence that every claim be demonstrable shaped the evidence-first character of "
    "this work.",
    "We are equally thankful to the Department of ICT and Computer Engineering, "
    f"{C.COLLEGE}, for providing the academic environment and the resources required to carry "
    "out this project, and to our teachers and friends for their reviews and criticism at "
    "every stage.",
    "We also acknowledge the open-source communities behind Node.js, React, Docker, nginx and "
    "the wider web-performance ecosystem, whose freely available tools made it possible to "
    "build and rigorously evaluate this project entirely at zero cost. Finally, we thank our "
    "families for their constant support."]:
    para(doc, t)
doc.add_paragraph()
for nm, rl in C.AUTHORS:
    para(doc, f"{nm} ({rl})", justify=False, after=0)

new_section(doc, "lowerRoman")
heading(doc, "ABSTRACT", 1)
for t in C.ABSTRACT:
    para(doc, t)

new_section(doc, "lowerRoman")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run_font(p.add_run("TABLE OF CONTENTS"), 16, bold=True)
add_toc(doc)

new_section(doc, "lowerRoman")
heading(doc, "LIST OF FIGURES", 1)
for b in C.BODY:
    if b["t"] == "fig":
        para(doc, b["cap"], justify=False, after=4)

new_section(doc, "lowerRoman")
heading(doc, "LIST OF TABLES", 1)
for b in C.BODY:
    if b["t"] == "table":
        para(doc, b["cap"], justify=False, after=4)

new_section(doc, "lowerRoman")
heading(doc, "LIST OF ACRONYMS / ABBREVIATIONS", 1)
ACR = [("ARE", "Adaptive Rendering Engine"), ("SSG", "Static Site Generation"),
       ("SSR", "Server-Side Rendering"), ("CSR", "Client-Side Rendering"),
       ("ISR", "Incremental Static Regeneration"),
       ("EDGE_ISR", "Edge Incremental Static Regeneration"),
       ("SWR", "Stale-While-Revalidate"), ("TTFB", "Time To First Byte"),
       ("FCP", "First Contentful Paint"), ("LCP", "Largest Contentful Paint"),
       ("TTL", "Time To Live"), ("HTTP", "HyperText Transfer Protocol"),
       ("HTML", "HyperText Markup Language"), ("JSON", "JavaScript Object Notation"),
       ("NDJSON", "Newline-Delimited JSON"), ("CDN", "Content Delivery Network"),
       ("API", "Application Programming Interface"),
       ("SEO", "Search Engine Optimization"), ("UA", "User-Agent"), ("TS", "TypeScript"),
       ("ab", "ApacheBench, the Apache HTTP server benchmarking tool")]
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
for a, b in ACR:
    c = t.add_row().cells
    set_cell(c[0], a, bold=True); set_cell(c[1], b)
    c[0].width = Inches(1.3); c[1].width = Inches(AVAIL_IN - 1.3)

# ---- body (arabic, bottom right, running header)
new_section(doc, "decimal", start=1, header=C.TITLE, align=WD_ALIGN_PARAGRAPH.RIGHT)

first_h1 = True
for blk in C.BODY:
    t = blk["t"]
    if t == "h1":
        heading(doc, blk["text"], 1, page_break=not first_h1)
        first_h1 = False
    elif t == "h2":
        heading(doc, blk["text"], 2)
    elif t == "h3":
        heading(doc, blk["text"], 3)
    elif t == "p":
        para(doc, blk["text"], italic=blk.get("italic", False))
    elif t == "quote":
        para(doc, blk["text"], italic=True, indent=0.35, before=4, after=10)
    elif t == "bul":
        add_list(doc, blk["items"], numbered=False)
    elif t == "num":
        add_list(doc, blk["items"], numbered=True)
    elif t == "refs":
        for i, it in enumerate(blk["items"], 1):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            pf.space_after = Pt(6)
            pf.left_indent = Inches(0.42); pf.first_line_indent = Inches(-0.42)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_font(p.add_run(f"[{i}]   {it}"), 12)
    elif t == "code":
        add_code(doc, blk["text"])
    elif t == "table":
        add_table(doc, blk)
    elif t == "fig":
        add_figure(doc, blk)
    elif t == "pb":
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

doc.save(OUT)
print(f"DOCX written: {OUT}")
print(f"   TOC entries seeded: {len(TOCMAP)}")
